"""Tests de parsers a CIR (Bloque 3)."""

import pytest
from docx import Document

from ai.errors import ScannedPDFError
from ai.tools.cir import ElementType
from ai.tools.parsers import DocxParser, TextToCIRAdapter
from ai.tools.parsers import pdf_parser as pdf_mod

# --- TextToCIRAdapter -------------------------------------------------------


def test_texto_estructurado():
    texto = (
        "# Proceso de Siniestros\n\n"
        "Este documento describe el registro de siniestros.\n\n"
        "- Reportar el siniestro\n- Registrar la guía\n- Cerrar\n"
    )
    cir = TextToCIRAdapter.adapt(texto, title="Siniestros")
    tipos = [e.type for e in cir.elements]
    assert ElementType.HEADING in tipos
    assert ElementType.LIST in tipos
    assert ElementType.PARAGRAPH in tipos
    # provenance/orden estables
    assert [e.element_id for e in cir.elements] == [
        f"el-{i:04d}" for i in range(len(cir.elements))
    ]


def test_texto_plano_seccion_es_rotulo_y_cuerpo_va_en_parrafo():
    """El texto de una SECTION es un RÓTULO, nunca el cuerpo del documento.

    ``CIRBuilder.add_section`` apila ese texto como ancestro del breadcrumb y el
    chunker lo usa como contexto del chunk, así que una sección que llevara el
    documento entero lo mandaría dos veces al modelo (2,00x medido). Es el
    invariante que el resto de parsers ya cumplía: todos los demás
    ``add_section`` del repositorio pasan un título.
    """
    texto = "solo una linea de texto plano sin estructura alguna aqui"
    cir = TextToCIRAdapter.adapt(texto)
    assert cir.source_type == "text"

    seccion, parrafo = cir.elements
    assert seccion.type is ElementType.SECTION
    assert seccion.text == "Documento"  # rótulo, no contenido
    assert parrafo.type is ElementType.PARAGRAPH
    assert parrafo.text == texto  # el contenido se conserva íntegro


def test_texto_plano_usa_el_titulo_dado_como_rotulo():
    cir = TextToCIRAdapter.adapt("texto plano cualquiera sin estructura", title="Guías")
    assert cir.elements[0].text == "Guías"


def test_documento_vacio_no_gana_un_rotulo_con_el_nombre_del_fichero():
    """Sin cuerpo no hay rótulo: se leería como contenido y sería citable."""
    cir = TextToCIRAdapter.adapt("   ", title="modernizacion.md")
    assert [e.text for e in cir.elements] == [""]


def test_ningun_parser_mete_el_cuerpo_en_una_seccion():
    """Candado: ninguna SECTION/HEADING lleva el documento como texto.

    Se comprueba sobre la propiedad, no sobre una longitud: el texto del
    elemento que abre un chunk NO puede ser el contenido del documento.
    """
    cuerpo = "El transportista registra la guia y actualiza el checkpoint. " * 400
    for texto in (cuerpo, "# Titulo\n\n" + cuerpo):
        cir = TextToCIRAdapter.adapt(texto)
        for el in cir.elements:
            if el.type in (ElementType.SECTION, ElementType.HEADING):
                assert el.text is not None
                assert cuerpo.strip() not in el.text


def test_breadcrumb_traza_headings():
    texto = "# A\n\n## B\n\nparrafo bajo B\n"
    cir = TextToCIRAdapter.adapt(texto, title="Doc")
    parrafo = next(e for e in cir.elements if e.type is ElementType.PARAGRAPH)
    # breadcrumb incluye la sección raíz y los headings ancestros
    assert "Doc" in parrafo.breadcrumb
    assert "B" in parrafo.breadcrumb


# --- DocxParser -------------------------------------------------------------


def test_docx_estructura_y_tabla_integra(tmp_path):
    doc = Document()
    doc.add_heading("Proceso de Siniestros", level=1)
    doc.add_paragraph("Introducción al proceso.")
    doc.add_paragraph("Reportar siniestro", style="List Bullet")
    doc.add_paragraph("Registrar guía", style="List Bullet")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Campo"
    table.cell(0, 1).text = "Tipo"
    table.cell(1, 0).text = "numero_guia"
    table.cell(1, 1).text = "texto"
    path = tmp_path / "proc.docx"
    doc.save(str(path))

    cir = DocxParser.parse(path)
    assert cir.fidelity == "full"
    assert any(e.type is ElementType.HEADING for e in cir.elements)
    assert any(e.type is ElementType.LIST for e in cir.elements)

    tablas = cir.tables()
    assert len(tablas) == 1
    tbl = tablas[0].table
    assert tbl.n_rows == 2 and tbl.n_cols == 2
    assert tbl.rows[0] == ["Campo", "Tipo"]
    assert tbl.rows[1] == ["numero_guia", "texto"]


# --- PdfParser (pypdf monkeypatcheado) --------------------------------------


class _FakePage:
    def __init__(self, text: str) -> None:
        self._text = text

    def extract_text(self) -> str:
        return self._text


def _fake_reader(pages_text):
    class _FakeReader:
        def __init__(self, _path):
            self.pages = [_FakePage(t) for t in pages_text]

    return _FakeReader


def test_pdf_texto_con_coordenadas_de_pagina(monkeypatch):
    monkeypatch.setattr(
        pdf_mod,
        "PdfReader",
        _fake_reader(["# Título\n\nUn párrafo de la página uno.", "Texto página dos."]),
    )
    cir = pdf_mod.PdfParser.parse("dummy.pdf")
    assert cir.fidelity == "degraded"
    assert any(e.type is ElementType.HEADING for e in cir.elements)
    assert any(e.coordinates.page == 2 for e in cir.elements)


def test_pdf_escaneado_lanza_error(monkeypatch):
    monkeypatch.setattr(pdf_mod, "PdfReader", _fake_reader(["", "   ", "\n"]))
    with pytest.raises(ScannedPDFError):
        pdf_mod.PdfParser.parse("scan.pdf")
