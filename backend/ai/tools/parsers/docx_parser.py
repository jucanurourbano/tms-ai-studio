"""DocxParser: convierte un .docx en un CIR (python-docx).

Recorre el cuerpo del documento en orden preservando párrafos, headings, listas
y tablas (con topología). Fidelidad: ``full``.
"""

from pathlib import Path
from typing import Iterator, Union

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from ai.errors import ParserError
from ai.tools.cir import CIR

from ._builder import CIRBuilder


def _iter_block_items(document: _Document) -> Iterator[Union[Paragraph, Table]]:
    """Itera párrafos y tablas en el orden real del documento."""
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _heading_level(style_name: str) -> Union[int, None]:
    """Devuelve el nivel si el estilo es Heading N / Title, si no None."""
    if not style_name:
        return None
    name = style_name.strip().lower()
    if name == "title":
        return 1
    if name.startswith("heading"):
        digits = name.replace("heading", "").strip()
        return int(digits) if digits.isdigit() else 1
    return None


class DocxParser:
    """Parser de documentos .docx."""

    @staticmethod
    def parse(path: Union[str, Path]) -> CIR:
        """Parsea un .docx a CIR."""
        path = Path(path)
        try:
            document = Document(str(path))
        except Exception as exc:  # pragma: no cover
            raise ParserError(f"No se pudo abrir el .docx: {exc}") from exc

        title = (document.core_properties.title or "").strip() or path.stem
        builder = CIRBuilder(source_type="document", fidelity="full", title=title)
        builder.add_section(title, level=0)

        list_buffer: list[str] = []

        def flush_list() -> None:
            nonlocal list_buffer
            if list_buffer:
                builder.add_list(list_buffer)
                list_buffer = []

        for block in _iter_block_items(document):
            if isinstance(block, Table):
                flush_list()
                rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
                builder.add_table(rows)
                continue

            text = block.text.strip()
            style_name = block.style.name if block.style else ""
            if not text:
                flush_list()
                continue

            level = _heading_level(style_name)
            if level is not None:
                flush_list()
                builder.add_heading(text, level=level)
            elif "list" in (style_name or "").lower():
                list_buffer.append(text)
            else:
                flush_list()
                builder.add_paragraph(text)

        flush_list()
        return builder.build()
